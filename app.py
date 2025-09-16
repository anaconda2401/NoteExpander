import google.generativeai as genai
from docx import Document
import gradio as gr
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import os


tesseract_path = "C:/Program Files/Tesseract-OCR/tesseract.exe" #Configure Tesseract path here
if os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

# 1. Configure Gemini API Key
genai.configure(api_key='Gemini-api-key')

# 2. Choose model (Flash = faster, Pro = stronger but quota limited)
model = genai.GenerativeModel("gemini-1.5-flash")

# 3. Extract text from PDF with OCR fallback (no Poppler needed)
def extract_text_from_pdf(pdf_file, ocr_threshold=50):
    full_text = ""
    doc = fitz.open(pdf_file.name)

    for i, page in enumerate(doc, start=1):
        text = page.get_text().strip()

        # If not enough text, render page as image + OCR
        if len(text) < ocr_threshold:
            print(f"⚠️ Page {i}: switching to OCR (found only {len(text)} chars)")
            pix = page.get_pixmap()  # render page as image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img)
            text += "\n[OCR Extracted]\n" + ocr_text

        full_text += f"\n--- Page {i} ---\n{text}\n"

    return full_text

# 4. Generate explained notes in single batch
def generate_explained_notes(full_text):
    prompt = f"""
🚨 STRICT INSTRUCTIONS 🚨
You are NOT allowed to summarize under any circumstances.
Your output MUST be at least as long as the input text, ideally 2x longer.
If the input is 70 pages, the output must be 70+ pages, NOT 8 pages.
If your response is shorter than the input, it is INVALID.

TASK:
Rewrite the given lecture notes into a fully expanded *textbook-style chapter*.

REQUIREMENTS:
- Expand every line into full teaching-style explanations.
- Cover every definition with plain-language explanations.
- Derive every formula step by step with intermediate steps.
- Add solved examples (easy, medium, hard).
- Insert analogies, real-world applications, and visual imagination aids.
- Discuss common mistakes and misconceptions for each concept.
- Turn short bullet points into multi-paragraph explanations.
- Write like a professor teaching students before an exam.

FORMATTING:
- Use Markdown with headings, subheadings, and numbered sections.
- Include "Worked Examples", "Common Mistakes", and "Extra Notes" sections where possible.

INPUT TO EXPAND:

{full_text}
"""
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            temperature=0.7,
            max_output_tokens=8192
        )
    )
    return response.text

# 5. Save explained notes into a Word file
def save_to_word(content, filename="Explained_Notes.docx"):
    doc = Document()
    doc.add_heading("Fully Explained Notes", 0)
    doc.add_paragraph(content)
    doc.save(filename)
    return filename

# 6. Main pipeline for Gradio
def process_pdf(pdf_file):
    raw_text = extract_text_from_pdf(pdf_file)
    print("⏳ Sending full document to Gemini...")
    explained_text = generate_explained_notes(raw_text)
    output_file = save_to_word(explained_text)
    return "✅ Notes generated successfully!", output_file

# 7. Gradio UI
with gr.Blocks() as demo:
    gr.Markdown("# 📘 Lecture Notes Expander (Gemini)\nUpload your PDF and get fully explained notes with scenarios & examples.")
    
    with gr.Row():
        pdf_input = gr.File(label="Upload your lecture PDF", file_types=[".pdf"])
    
    generate_btn = gr.Button("Generate Explained Notes (Single Batch)")
    status_output = gr.Textbox(label="Status")
    file_output = gr.File(label="Download Explained Notes (.docx)")
    
    generate_btn.click(process_pdf, inputs=pdf_input, outputs=[status_output, file_output])

# Run local server
if __name__ == "__main__":
    demo.launch(server_name="127.0.0.1", server_port=7860)

